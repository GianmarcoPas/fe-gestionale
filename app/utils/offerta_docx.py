from __future__ import annotations

from copy import deepcopy
from datetime import date
from io import BytesIO
from pathlib import Path
import re
from typing import Any

from docx import Document


MONTHS_IT_FULL = [
    "",
    "gennaio",
    "febbraio",
    "marzo",
    "aprile",
    "maggio",
    "giugno",
    "luglio",
    "agosto",
    "settembre",
    "ottobre",
    "novembre",
    "dicembre",
]


def format_eur(amount: float | int) -> str:
    try:
        val = float(amount or 0)
    except Exception:
        val = 0.0
    s = "{:,.2f}".format(val).replace(",", "X").replace(".", ",").replace("X", ".")
    return s


def format_date_it_long(d: date) -> str:
    return f"{d.day} {MONTHS_IT_FULL[d.month]} {d.year}"


def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    """
    Sostituisce placeholder preservando la formattazione dei RUN del placeholder.
    Supporta placeholder spezzati su più run (caso tipico nei .docx).
    """
    if not getattr(paragraph, "runs", None):
        return

    # Loop finché troviamo sostituzioni (per gestire più placeholder nello stesso paragrafo)
    while True:
        full = "".join(r.text for r in paragraph.runs)
        found = None  # (start, end, key, replacement)
        for k, v in mapping.items():
            idx = full.find(k)
            if idx != -1:
                found = (idx, idx + len(k), k, v)
                break
        if not found:
            break

        start, end, _k, repl = found

        # Mappa posizioni globali -> run index + offset
        cur = 0
        s_run = s_off = e_run = e_off = None
        for i, run in enumerate(paragraph.runs):
            nxt = cur + len(run.text)
            if s_run is None and start < nxt:
                s_run = i
                s_off = start - cur
            if e_run is None and end <= nxt:
                e_run = i
                e_off = end - cur
                break
            cur = nxt

        if s_run is None or e_run is None:
            break

        if s_run == e_run:
            run = paragraph.runs[s_run]
            run.text = run.text[:s_off] + repl + run.text[e_off:]
        else:
            start_run = paragraph.runs[s_run]
            end_run = paragraph.runs[e_run]

            # Mantieni lo stile: il testo sostituito resta nel run iniziale (stile placeholder)
            start_run.text = start_run.text[:s_off] + repl
            # Il resto del testo dopo il placeholder resta nel run finale (stile originale)
            end_run.text = end_run.text[e_off:]

            # Svuota i run intermedi (parte del placeholder)
            for j in range(s_run + 1, e_run):
                paragraph.runs[j].text = ""


def _replace_in_cell(cell, mapping: dict[str, str]) -> None:
    for p in cell.paragraphs:
        _replace_in_paragraph(p, mapping)
    # Tabelle annidate dentro la cella (Word permette nested tables)
    for t in getattr(cell, "tables", []):
        for row in t.rows:
            for c in row.cells:
                _replace_in_cell(c, mapping)


def _replace_everywhere(doc: Document, mapping: dict[str, str]) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                _replace_in_cell(cell, mapping)

    # Header / footer
    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_in_paragraph(p, mapping)
        for t in section.header.tables:
            for row in t.rows:
                for cell in row.cells:
                    _replace_in_cell(cell, mapping)

        for p in section.footer.paragraphs:
            _replace_in_paragraph(p, mapping)
        for t in section.footer.tables:
            for row in t.rows:
                for cell in row.cells:
                    _replace_in_cell(cell, mapping)


def _iter_all_tables(doc: Document):
    """Itera tutte le tabelle, incluse quelle annidate dentro le celle."""
    queue = list(doc.tables)
    while queue:
        t = queue.pop(0)
        yield t
        for row in t.rows:
            for cell in row.cells:
                for nt in getattr(cell, "tables", []):
                    queue.append(nt)


def _find_beni_table_and_row(doc: Document) -> tuple[Any, int] | tuple[None, None]:
    marker = "[bene 1]"
    for table in _iter_all_tables(doc):
        for ridx, row in enumerate(table.rows):
            for cell in row.cells:
                if marker in (cell.text or "").lower():
                    return table, ridx
    return None, None


def _populate_beni_table(doc: Document, beni: list[dict[str, Any]]) -> None:
    """
    Cerca la riga che contiene [bene 1] e la usa come template.
    Se ci sono N beni, duplica la riga N-1 volte e sostituisce i placeholder.
    """
    table, row_idx = _find_beni_table_and_row(doc)
    if table is None or row_idx is None:
        # Nessuna tabella beni trovata: niente da fare
        return

    template_row = table.rows[row_idx]
    # IMPORTANT: salva una copia "pulita" della riga template PRIMA di compilarla,
    # altrimenti i cloni successivi erediteranno già i valori del Bene 1.
    template_tr = deepcopy(template_row._tr)

    ph_re = re.compile(r"\[\s*(bene|valore\s*bene|importo\s*offerta)\s*1\s*\]", re.IGNORECASE)

    def row_has_placeholder(row) -> bool:
        for cell in row.cells:
            txt = (cell.text or "").replace("\u00A0", " ").replace("\u202F", " ")
            if ph_re.search(txt):
                return True
        return False

    def fill_row(row, bene: dict[str, Any]) -> None:
        desc = str(bene.get("descrizione") or "")
        val = format_eur(bene.get("valore") or 0)
        imp = format_eur(bene.get("importo_offerta") or 0)

        # Supporta placeholder con maiuscole/minuscole (es. [Bene 1])
        # Supporta anche NBSP e spazi doppi (Word spesso usa caratteri non visibili).
        def variants(s: str) -> set[str]:
            vs = {s}
            vs.add(s.replace(" ", "\u00A0"))      # NBSP
            vs.add(s.replace(" ", "\u202F"))      # narrow no-break space
            vs.add(s.replace(" ", "  "))          # double space
            return vs

        bene_keys = variants("[bene 1]") | variants("[Bene 1]") | variants("[BENE 1]")
        val_keys = variants("[valore bene 1]") | variants("[Valore bene 1]") | variants("[VALORE BENE 1]")
        imp_keys = variants("[importo offerta 1]") | variants("[Importo offerta 1]") | variants("[IMPORTO OFFERTA 1]")

        mapping = {
            **{k: desc for k in bene_keys},
            **{k: val for k in val_keys},
            **{k: imp for k in imp_keys},
        }
        for cell in row.cells:
            _replace_in_cell(cell, mapping)

        # Fallback ultra-robusto: se dopo la sostituzione restano placeholder,
        # forza il testo per colonne (può perdere lo stile, ma evita righe vuote).
        if row_has_placeholder(row) and len(row.cells) >= 3:
            row.cells[0].text = desc
            row.cells[1].text = val
            row.cells[2].text = imp

    # Compila prima riga, poi duplica inserendo DOPO la riga precedente.
    # Nota: usare `addnext` è più affidabile di `_tbl.insert(...)` (che può inserirsi in punti non-row).
    prev_tr = template_row._tr
    for i, bene in enumerate(beni):
        if i == 0:
            fill_row(template_row, bene)
            prev_tr = template_row._tr
            continue

        new_tr = deepcopy(template_tr)
        prev_tr.addnext(new_tr)

        # Recupera l'indice reale della nuova riga appena inserita
        idx = list(table._tbl.tr_lst).index(new_tr)
        new_row = table.rows[idx]
        fill_row(new_row, bene)
        prev_tr = new_tr

    # Cleanup: rimuove eventuali righe "di esempio" rimaste con placeholder (es. una riga sopra l'header)
    # che alcuni template possono contenere.
    # rimuovi dal basso verso l'alto per non rompere gli indici
    for ridx in range(len(table.rows) - 1, -1, -1):
        row = table.rows[ridx]
        if row_has_placeholder(row):
            table._tbl.remove(row._tr)


def generate_offerta_docx(
    template_path: Path,
    *,
    cliente_nome: str,
    indirizzo: str,
    civico: str,
    cap: str,
    comune: str,
    prov: str,
    piva: str,
    data_emissione_text: str,
    importo_caricamento: float,
    importo_revisione: float,
    beni: list[dict[str, Any]],
) -> BytesIO:
    doc = Document(str(template_path))

    mapping = {
        "[Cliente]": cliente_nome or "",
        "[Indirizzo]": indirizzo or "",
        "[Civico]": civico or "",
        "[Cap]": cap or "",
        "[Comune]": comune or "",
        "[Prov]": prov or "",
        "[P. IVA]": piva or "",
        "[data emissione]": data_emissione_text or "",
        "[importo caricamento]": format_eur(importo_caricamento),
        "[importo revisione]": format_eur(importo_revisione),
    }

    _replace_everywhere(doc, mapping)
    _populate_beni_table(doc, beni)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

